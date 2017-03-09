# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: http://doc.scrapy.org/en/latest/topics/item-pipeline.html
from docx import *

class LiyuxiangPipeline(object):
    def __init__(self):
        if r"..\目录.docx":
            self.file = Document("..\目录.docx")
        else:
            self.file = Document()

    def process_item(self, items, spider):
        titles = items["title"]
        urls = items["url"]
        for i in len(titles):
            self.file.add_paragraph(titles[i])
            self.file.add_paragraph(urls[i])
        return items

    def open_spider(self):
        pass

    def close_spider(self):
        self.file.save("..\目录.docx")